from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
import re
import io

_FOOTNOTE_REF_PATTERN = re.compile(r'\[\[FOOTNOTE_REF:(\d+)\]\]')

_TOKEN_PATTERN = re.compile(
    r'(\[\[FOOTNOTE_REF:\d+\]\]'
    r'|\[\[ENTITY_REF:[^:]+:[^\]]+\]\]'
    r'|\[\[ENTITY_LINK:[^:]+:[^:]+:[^:]+:[^\]]*\]\]'
    r'|\[\[TWIST_REF:[^:]+:[^\]]+\]\])'
)

_ENTITY_STYLES = {
    'char':   (True,  False, None),
    'loc':    (False, True,  None),
    'item':   (False, True,  None),
    'lore':   (False, True,  None),
    'group':  (True,  False, None),
    'secret': (False, False, '990000'),
}

# top/bottom margins per trim size (industry standard)
_TRIM_MARGINS = {
    'pocket':   (0.60, 0.70),
    'standard': (0.75, 0.85),
    'large':    (0.75, 0.85),
}

def render(project_title, author_name, chapters, content_mode, overrides=None) -> bytes:
    """
    Renders chapters to a professional DOCX document.
    Standard Manuscript Format (Double spaced, TNR 12pt, 1" margins)
    or Book-Ready format based on overrides.
    Returns bytes of the generated file.
    """
    doc = Document()
    
    # 0. Formatting Constants
    # TRIM_SIZES: { key: (width_inches, height_inches) }
    TRIM_SIZES = {
        'pocket': (4.25, 6.87),
        'standard': (5.0, 8.0),
        'large': (6.0, 9.0)
    }
    
    # Defaults
    trim_key = overrides.get('trim', 'standard') if overrides else 'standard'
    font_size = overrides.get('font_size', 12) if overrides else 12
    # gutter/outer come pre-computed from the frontend (industry-accurate values)
    gutter = overrides.get('gutter', 0.5) if overrides else 0.5
    outer  = overrides.get('outer',  0.4) if overrides else 0.4
    if gutter is None: gutter = 0.5
    if outer  is None: outer  = 0.4

    width, height = TRIM_SIZES.get(trim_key, TRIM_SIZES['standard'])
    top_in, bottom_in = _TRIM_MARGINS.get(trim_key, (0.75, 0.85))

    # 1. Setup Document Sections (Margins and Size)
    section = doc.sections[0]
    section.page_width = Inches(width)
    section.page_height = Inches(height)

    # Inner (gutter) margin on left, outer margin on right.
    # Mirror margins are enabled below so these alternate correctly for print.
    section.top_margin    = Inches(top_in)
    section.bottom_margin = Inches(bottom_in)
    section.left_margin   = Inches(gutter)
    section.right_margin  = Inches(outer)

    # Enable mirror margins so the gutter alternates to the binding edge on each page
    section_props = section._sectPr
    mirror = OxmlElement('w:mirrorMargins')
    section_props.append(mirror)
    
    # 2. Setup Default Style (Normal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(font_size)
    
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.first_line_indent = Inches(0.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # 3. Title Page (Optional but good for Stage 2)
    # For now, let's just start with the content to keep it simple and clean.
    
    # 4. Chapters
    for idx, chapter in enumerate(chapters):
        # Chapter Heading
        # We don't want indent for headings
        h = doc.add_paragraph()
        h_run = h.add_run(chapter.get('title', f"Chapter {idx+1}").upper())
        h_run.bold = True
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.first_line_indent = Inches(0)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
        
        text = chapter.get('text', '')
        footnotes = chapter.get('footnotes', [])

        text = re.sub(r'\{(secret|knows|believes):([^}]+)\}', r'', text)

        def _add_styled_paragraph(doc, raw_text):
            """Split raw_text on tokens and build a paragraph with mixed runs."""
            tokens = _TOKEN_PATTERN.split(raw_text)
            if len(tokens) == 1 and not _TOKEN_PATTERN.match(raw_text):
                # No tokens — plain paragraph
                doc.add_paragraph(raw_text)
                return

            p = doc.add_paragraph()
            for tok in tokens:
                if not tok:
                    continue

                fn_m = re.fullmatch(r'\[\[FOOTNOTE_REF:(\d+)\]\]', tok)
                if fn_m:
                    ref_run = p.add_run(f"[{fn_m.group(1)}]")
                    ref_run.font.superscript = True
                    ref_run.font.bold = True
                    ref_run.font.size = Pt(max(7, font_size - 3))
                    continue

                er_m = re.fullmatch(r'\[\[ENTITY_REF:([^:]+):([^\]]+)\]\]', tok)
                if er_m:
                    etype, name = er_m.group(1), er_m.group(2)
                    bold, italic, color = _ENTITY_STYLES.get(etype, (False, False, None))
                    run = p.add_run(name)
                    run.bold = bold
                    run.italic = italic
                    if color:
                        run.font.color.rgb = RGBColor.from_string(color)
                    continue

                el_m = re.fullmatch(r'\[\[ENTITY_LINK:([^:]+):([^:]+):([^:]+):([^\]]*)\]\]', tok)
                if el_m:
                    etype, eid, name, desc = el_m.group(1), el_m.group(2), el_m.group(3), el_m.group(4)
                    bold, italic, color = _ENTITY_STYLES.get(etype, (False, True, None))
                    run = p.add_run(name)
                    run.bold = bold
                    run.italic = italic
                    if color:
                        run.font.color.rgb = RGBColor.from_string(color)
                    continue

                tw_m = re.fullmatch(r'\[\[TWIST_REF:([^:]+):([^\]]+)\]\]', tok)
                if tw_m:
                    ttype, name = tw_m.group(1), tw_m.group(2)
                    run = p.add_run(name)
                    if ttype == 'twist':
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # dark green
                    else:  # foreshadow
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x78, 0x55, 0x9A)  # muted purple
                    continue

                # Plain text segment
                p.add_run(tok)

        # Process paragraphs
        paragraphs = text.split('\n')
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue

            if p_text in ('---', '***', '* * *'):
                sb = doc.add_paragraph("* * *")
                sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sb.paragraph_format.first_line_indent = Inches(0)
            else:
                _add_styled_paragraph(doc, p_text)

        # Footnotes section at bottom of chapter (before page break)
        if footnotes:
            sep = doc.add_paragraph()
            sep.add_run("─" * 32)
            sep.paragraph_format.first_line_indent = Inches(0)
            sep.paragraph_format.space_before = Pt(6)
            sep.paragraph_format.space_after = Pt(4)

            for fn_idx, fn_text in enumerate(footnotes):
                fn_p = doc.add_paragraph()
                fn_p.paragraph_format.first_line_indent = Inches(0)
                fn_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                fn_p.paragraph_format.space_before = Pt(0)
                fn_p.paragraph_format.space_after = Pt(2)

                num_run = fn_p.add_run(f"[{fn_idx + 1}]  ")
                num_run.font.superscript = True
                num_run.font.bold = True
                num_run.font.size = Pt(max(7, font_size - 3))

                content_run = fn_p.add_run(fn_text or "")
                content_run.font.size = Pt(max(8, font_size - 2))

        # Chapter Break (Page Break)
        if idx < len(chapters) - 1:
            doc.add_page_break()

    # Save to BytesIO
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()
